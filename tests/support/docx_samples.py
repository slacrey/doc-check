from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Mm, Pt


def ensure_docx_samples(fixtures_dir: Path) -> dict[str, Path]:
    fixtures_dir.mkdir(parents=True, exist_ok=True)
    valid_path = fixtures_dir / "aeos-valid-sample.docx"
    format_errors_path = fixtures_dir / "aeos-format-errors.docx"
    news_valid_path = fixtures_dir / "news-valid-sample.docx"
    news_errors_path = fixtures_dir / "news-format-errors.docx"
    aeos_standard_path = fixtures_dir / "aeos-standard-rules.docx"
    speech_valid_path = fixtures_dir / "speech-valid-sample.docx"
    speech_errors_path = fixtures_dir / "speech-format-errors.docx"
    speech_salutation_errors_path = fixtures_dir / "speech-salutation-errors.docx"
    speech_slogan_errors_path = fixtures_dir / "speech-slogan-errors.docx"

    _build_valid_sample(valid_path)
    _build_format_error_sample(format_errors_path)
    _build_news_valid_sample(news_valid_path)
    _build_news_error_sample(news_errors_path)
    _build_aeos_standard_sample(aeos_standard_path)
    _build_speech_valid_sample(speech_valid_path)
    _build_speech_error_sample(speech_errors_path)
    _build_speech_salutation_error_sample(speech_salutation_errors_path)
    _build_speech_slogan_error_sample(speech_slogan_errors_path)

    return {
        "valid": valid_path,
        "format_errors": format_errors_path,
        "news_valid": news_valid_path,
        "news_errors": news_errors_path,
        "aeos_standard": aeos_standard_path,
        "speech_valid": speech_valid_path,
        "speech_errors": speech_errors_path,
        "speech_salutation_errors": speech_salutation_errors_path,
        "speech_slogan_errors": speech_slogan_errors_path,
    }


def _build_valid_sample(path: Path) -> None:
    document = Document()
    _apply_gbt_9704_layout(document)
    body_style = _ensure_paragraph_style(
        document,
        "AEOS Body",
        font_name="仿宋_GB2312",
        font_size_pt=16,
    )
    toc_one = _ensure_paragraph_style(document, "TOC 1", base_style_name="Normal")
    toc_two = _ensure_paragraph_style(document, "TOC 2", base_style_name="Normal")

    document.add_paragraph("1 总则\t1", style=toc_one.name)
    document.add_paragraph("1.1 目的\t2", style=toc_two.name)
    document.add_paragraph("1 总则", style="Heading 1")
    document.add_paragraph("本制度适用于 AEOS 管理体系。", style=body_style.name)
    document.add_paragraph("网络与信息安全要求统一执行。", style=body_style.name)

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格中的适用条款"
    table.cell(0, 0).paragraphs[0].style = body_style

    document.add_page_break()
    document.add_paragraph("1.1 目的", style="Heading 2")
    document.add_paragraph("统一检查格式和术语。", style=body_style.name)

    section = document.sections[0]
    section.header.paragraphs[0].text = "AEOS 文件页眉"
    section.footer.paragraphs[0].text = "AEOS 文件页脚"

    document.save(path)


def _build_format_error_sample(path: Path) -> None:
    document = Document()
    body_style = _ensure_paragraph_style(
        document,
        "AEOS Body",
        font_name="SimSun",
        font_size_pt=12,
    )

    document.add_paragraph("1 总则", style="Heading 1")
    inherited = document.add_paragraph("继承样式的正文段落。", style=body_style.name)
    inherited.paragraph_format.alignment = None

    nonstandard = document.add_paragraph("直接设置为单倍行距的段落。", style=body_style.name)
    nonstandard.paragraph_format.line_spacing = 1.0
    nonstandard.paragraph_format.space_after = Pt(0)
    document.add_paragraph("网络安全,黑名单管理要求。", style=body_style.name)

    section = document.sections[0]
    section.header.paragraphs[0].text = "格式错误样本页眉"
    section.footer.paragraphs[0].text = "格式错误样本页脚"

    document.save(path)


def _build_news_valid_sample(path: Path) -> None:
    document = Document()
    document.add_paragraph("公司举办数据安全专题培训")
    document.add_paragraph("来源：管理与信息化部  2026年4月22日")
    document.add_paragraph("2026年4月22日，公司举办数据安全专题培训，系统讲解制度执行要点。")
    document.add_paragraph("培训围绕数据安全责任落实、制度宣贯和问题整改展开。")
    document.save(path)


def _build_news_error_sample(path: Path) -> None:
    document = Document()
    document.add_paragraph("史上最强！公司培训炸裂来袭！全员直呼太震撼了！")
    document.add_paragraph("2026/04/22，公司开展活动，网传效果非常震撼，实现绝对安全。")
    document.add_paragraph("活动现场气氛热烈，大家反响强烈。")
    document.save(path)


def _build_aeos_standard_sample(path: Path) -> None:
    document = Document()
    _apply_gbt_9704_layout(document)
    body_style = _ensure_paragraph_style(
        document,
        "AEOS Standard Body",
        font_name="仿宋_GB2312",
        font_size_pt=16,
    )

    document.add_paragraph("AEOS 制度文件编写规范", style="Heading 1")
    document.add_paragraph("正文统一使用“网络与信息安全”，不要使用“网络安全”。", style=body_style.name)
    document.add_paragraph("制度文件中不得使用“黑名单”表述。", style=body_style.name)
    document.save(path)


def _build_speech_valid_sample(path: Path) -> None:
    document = Document()
    document.add_paragraph("在2026年工作推进会上的讲话")
    document.add_paragraph("各位领导、同志们：")
    document.add_paragraph("今天，我们召开工作推进会，目的是进一步统一思想、压实责任。")
    document.add_paragraph("下面，我讲三点意见。")
    document.add_paragraph("同志们，让我们真抓实干，确保各项任务落地见效。谢谢大家！")
    document.save(path)


def _build_speech_error_sample(path: Path) -> None:
    document = Document()
    document.add_paragraph("关于推进工作的说明材料")
    document.add_paragraph("来源：办公室")
    document.add_paragraph("记者讯，今天我们简单说一下有关情况。")
    document.add_paragraph("公司表示，这项工作绝绝子，效果非常炸裂。")
    document.save(path)


def _build_speech_salutation_error_sample(path: Path) -> None:
    document = Document()
    document.add_paragraph("在2026年工作推进会上的讲话")
    document.add_paragraph("各位领导、同志们")
    document.add_paragraph("今天，我们召开工作推进会，目的是统一思想、压实责任。")
    document.add_paragraph("下面，我讲三点意见。")
    document.add_paragraph("谢谢大家。")
    document.save(path)


def _build_speech_slogan_error_sample(path: Path) -> None:
    document = Document()
    document.add_paragraph("在2026年工作推进会上的讲话")
    document.add_paragraph("各位领导、同志们：")
    document.add_paragraph("今天，我们召开工作推进会，目的是统一思想、压实责任。")
    document.add_paragraph("下面，我讲三点意见。")
    document.add_paragraph("最后，统一思想、凝聚力量、攻坚克难、狠抓落实！谢谢大家。")
    document.save(path)


def _ensure_paragraph_style(
    document: Document,
    name: str,
    *,
    base_style_name: str = "Normal",
    font_name: str = "SimSun",
    font_size_pt: float = 12,
):
    try:
        style = document.styles[name]
    except KeyError:
        style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    style.base_style = document.styles[base_style_name]
    style.font.name = font_name
    style.font.size = Pt(font_size_pt)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.left_indent = Pt(24)
    style.paragraph_format.space_after = Pt(6)
    return style


def _apply_gbt_9704_layout(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)
